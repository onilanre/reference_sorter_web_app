from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FILE = 'sorted_references.docx'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/sort', methods=['POST'])
def sort_references():
    uploaded_file = request.files.get('docx_file')
    if uploaded_file and uploaded_file.filename.endswith('.docx'):
        input_path = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
        uploaded_file.save(input_path)

        doc = Document(input_path)
        references = [para.text for para in doc.paragraphs if para.text.strip() != '']
        sorted_refs = sorted(references, key=lambda x: x.lower())

        sorted_doc = Document()
        for ref in sorted_refs:
            sorted_doc.add_paragraph(ref)
        sorted_doc.save(OUTPUT_FILE)

        return send_file(OUTPUT_FILE, as_attachment=True)
    else:
        return "Only .docx files are supported.", 400

if __name__ == '__main__':
    app.run(debug=True)
